import os
import re
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO
from docx.shared import Pt
from textwrap import wrap

# Function to process data and create tables
def process_data(data):
    tabel_per_matakuliah = {}

    # Drop duplicate entries
    data = data.drop_duplicates(subset=['Mata Kuliah', 'Pertanyaan'])

    for mata_kuliah in data['Mata Kuliah'].unique():
        subset_tabel = data[data['Mata Kuliah'] == mata_kuliah].copy()

        # Convert percentage columns to numeric
        subset_tabel[['Sangat Setuju', 'Setuju', 'Tidak Setuju', 'Sangat Tidak Setuju']] = subset_tabel[['Sangat Setuju', 'Setuju', 'Tidak Setuju', 'Sangat Tidak Setuju']].replace('%', '', regex=True).astype(float)

        # Setelah diubah ke numerik, kita bisa mengganti nilai '%' di sini
        subset_tabel = subset_tabel[['Pertanyaan', 'Sangat Setuju', 'Setuju', 'Tidak Setuju', 'Sangat Tidak Setuju']].set_index('Pertanyaan')

        # Format values to have two decimal places
        subset_tabel = subset_tabel.round(2)

        tabel_per_matakuliah[mata_kuliah] = subset_tabel

    return tabel_per_matakuliah

# Function to create pie chart for each pertanyaan
def create_pie_chart_per_pertanyaan(data, mata_kuliah):
    figures = []  # Store figures to be displayed later

    for idx, (pertanyaan, values) in enumerate(data.iterrows(), start=1):
        fig, ax = plt.subplots(figsize=(8, 6))  # Adjust the figsize as needed

        labels = data.columns
        sizes = values.values
        colors = ['gold', 'yellowgreen', 'lightcoral', 'lightskyblue']
        explode = (0.1, 0, 0, 0)  # explode 1st slice

        # Check if all values are zero
        if (sizes == 0).all():
            sizes = [1] * len(sizes)  # Set all sizes to 1 if they are all zero

        try:
            wrapped_title = "\n".join(wrap(f"{idx}. {pertanyaan}", 55))
            title_with_newline = wrapped_title + '\n'  # Add a newline after the title
            ax.set_title(title_with_newline, fontsize=16, fontweight='bold')

            wedges, texts, autotexts = ax.pie(sizes, explode=explode, labels=labels, colors=colors,
                                              autopct=lambda p: '{:.0f} ({:.1f}%)'.format(p * sum(sizes) / 100, p),
                                              textprops={'fontsize': 16},  # Set the font size for the text in the pie chart
                                              shadow=True, startangle=140)
            ax.axis('equal')

            figures.append(fig)
        except Exception as e:
            print(f"Error: {e}")

    return figures

def save_figures_to_word(figures, document):
    num_figures = len(figures)
    figures_per_page = 6  # Number of figures to display per page (2 pie charts per page)
    num_pages = -(-num_figures // figures_per_page)  # Calculate the number of pages needed

    for page_num in range(num_pages):
        # Add a page break for all pages except the first
        if page_num > 0:
            document.add_page_break()

        # Calculate the start and end index for the figures on this page
        start_idx = page_num * figures_per_page
        end_idx = min((page_num + 1) * figures_per_page, num_figures)

        # Create a table to hold the figures (1 row, 2 columns for 2 pie charts)
        table = document.add_table(rows=3, cols=2)

        for idx, fig_idx in enumerate(range(start_idx, end_idx)):
            cell = table.cell(0, idx)
            
            # Add the figure to the cell
            image_stream = BytesIO()
            figures[fig_idx].savefig(image_stream, format='png')
            plt.close(figures[fig_idx])
            cell.paragraphs[0].add_run().add_picture(image_stream, width=Inches(3), height=Inches(2.5))

def format_lecturer_name(name):
    # Use regular expression to insert spaces before capital letters
    formatted_name = re.sub(r"([A-Z])", r" \1", name).strip()
    return formatted_name

# Main function for generating Word document
def generate_word_document(data, path_data):
    # Extract lecturer's name from file name
    lecturer_name_match = re.search(r'_([A-Za-z]+)STMT', path_data)
    lecturer_name = lecturer_name_match.group(1) if lecturer_name_match else "Unknown"

    # Format lecturer's name with spaces
    formatted_lecturer_name = format_lecturer_name(lecturer_name)

    # Process data
    tabel_per_matakuliah = process_data(data)

    # Create a Word document
    document = Document()

    first_page = True  # Flag to identify the first page

    # Display tables and create pie chart
    for mata_kuliah, tabel in tabel_per_matakuliah.items():

        if not first_page:
            document.add_page_break()
        else:
            first_page = False  # Set the flag to False after the first table

        document.add_heading(f"{mata_kuliah}", level=1,)

        # Add lecturer name and number of respondents in a single paragraph
        lecturer_and_respondents_paragraph = document.add_paragraph()

        # Add lecturer name on the left
        lecturer_and_respondents_paragraph.add_run(f"Dosen: {formatted_lecturer_name}").font.size = Pt(12)

        # Add spaces for alignment
        lecturer_and_respondents_paragraph.add_run().add_text(" " * 10)

        # Add number of respondents on the right
        lecturer_and_respondents_paragraph.add_run(f"Jumlah Responden: ").font.size = Pt(12)

        # Add table to Word document
        table = document.add_table(rows=tabel.shape[0] + 1, cols=tabel.shape[1] + 2)  # Add 2 columns for 'Nilai'

        # Adjust column widths
        for col_num, column in enumerate(table.columns):
            if col_num == 0:
                table.cell(0, col_num).width = Inches(10)  # Adjust the width for Pertanyaan column
            elif col_num == len(table.columns) - 1 or col_num == len(table.columns):  # Set the width for 'Nilai' columns
                table.cell(0, col_num).width = Inches(0.5)
            else:
                table.cell(0, col_num).width = Inches(0.5)  # Adjust the width for other columns

        # Add column headers (Pertanyaan, values, and 'Nilai')
        table.cell(0, 0).text = 'Pertanyaan'
        table.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(12)
        table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        for col_num, column_name in enumerate(tabel.columns):
            cell = table.cell(0, col_num + 1)
            cell.text = column_name
            cell.paragraphs[0].runs[0].font.bold = True  # Make the header bold
            cell.paragraphs[0].runs[0].font.size = Pt(12)  # Adjust the font size for the header

        # Add 'Nilai' column header
        table.cell(0, len(table.columns) - 1).text = 'Nilai'
        table.cell(0, len(table.columns) - 1).paragraphs[0].runs[0].font.size = Pt(12)
        table.cell(0, len(table.columns) - 1).paragraphs[0].runs[0].font.bold = True

        for row_num, (pertanyaan, values) in enumerate(tabel.iterrows(), start=1):
            table.cell(row_num, 0).text = f"{row_num}. {pertanyaan}"  # Set the combined "No" and "Pertanyaan" column
            table.cell(row_num, 0).paragraphs[0].runs[0].font.size = Pt(12)  # Set the font size for the combined column
            for col_num, value in enumerate(values, start=1):
                cell = table.cell(row_num, col_num)

                if col_num != len(table.columns):  # Exclude the 'Nilai' column
                    # Add a percentage sign to each value
                    cell.text = "{:.2f}%".format(value)
                else:
                    # Divide the 'Nilai' column values by 100 and add a percentage sign
                    cell.text = "{:.2f}%".format(value / 100)

                cell.paragraphs[0].runs[0].font.size = Pt(12)

            # Calculate the 'Nilai' column value based on the provided formula for each row
            nilai_value = values['Sangat Setuju'] * 4 + values['Setuju'] * 3 + values['Tidak Setuju'] * 2 + values['Sangat Tidak Setuju'] * 1
            # Set the 'Nilai' column value for each row
            table.cell(row_num, len(table.columns) - 1).text = "{:.2f}".format(nilai_value/100)
            table.cell(row_num, len(table.columns) - 1).paragraphs[0].runs[0].font.size = Pt(12)

        # Add heading for pie charts
        document.add_page_break()
        # Create pie chart for each pertanyaan on a new page
        figures = create_pie_chart_per_pertanyaan(tabel, mata_kuliah)

        # Save figures to Word document
        save_figures_to_word(figures, document)

    # Set line spacing and space after
    # Save the Word document
    document.save(f"KL_Kuesioner_202312_{lecturer_name}_STMT (1).docx")

if __name__ == '__main__':
    # Read data CSV
    current_directory = os.path.dirname(os.path.abspath(__file__))
    path_data = os.path.join(current_directory, 'KL_Kuesioner_202312_MochammadFathurridhoHermantoSTMT (1).csv')
    data = pd.read_csv(path_data)

    # Generate Word document
    generate_word_document(data, path_data)
